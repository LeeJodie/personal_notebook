"""Isolated DOCX, PDF, XLSX and Markdown reader-model processor.

The worker sends the private R2 object stream to this service. The service never
persists uploads and emits only the normalized reader model used by H5, TTS and
knowledge indexing.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import unquote

import fitz
from docx import Document
from fastapi import FastAPI, HTTPException, Request
from openpyxl import load_workbook

app = FastAPI(title="ShengYue document processor", version="1.0.0")

MAX_BYTES = 200 * 1024 * 1024
MAX_SECTIONS = 100
MAX_PARAGRAPHS_PER_SECTION = 32
MAX_TEXT_CHARS = 120_000


def clean(value: Any) -> str:
    return " ".join(str(value or "").replace("\x00", " ").split())


def section(title: str, paragraphs: list[str], index: int) -> dict[str, Any] | None:
    usable = [clean(paragraph)[:4_000] for paragraph in paragraphs if clean(paragraph)]
    if not usable:
        return None
    return {
        "id": f"section-{index + 1}",
        "eyebrow": f"{index + 1:02d} / 文档内容",
        "title": clean(title)[:300] or f"第 {index + 1} 节",
        "paragraphs": usable[:MAX_PARAGRAPHS_PER_SECTION],
    }


def markdown_sections(text: str) -> list[dict[str, Any]]:
    groups: list[tuple[str, list[str]]] = []
    title, paragraphs, buffer = "正文", [], []

    def flush_paragraph() -> None:
        nonlocal buffer
        content = clean(" ".join(buffer))
        if content:
            paragraphs.append(content)
        buffer = []

    def flush_section() -> None:
        flush_paragraph()
        if paragraphs:
            groups.append((title, paragraphs.copy()))

    for raw_line in text.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if line.startswith("#") and line.lstrip("#").startswith(" "):
            flush_section()
            title = clean(line.lstrip("#"))
            paragraphs = []
        elif not line:
            flush_paragraph()
        else:
            buffer.append(line.lstrip("-*+ "))
    flush_section()
    return [candidate for index, (heading, lines) in enumerate(groups[:MAX_SECTIONS]) if (candidate := section(heading, lines, index))]


def parse_docx(blob: bytes) -> tuple[str, list[dict[str, Any]]]:
    document = Document(BytesIO(blob))
    title = clean(document.core_properties.title)
    groups: list[tuple[str, list[str]]] = []
    current_title, current_paragraphs = "正文", []
    for paragraph in document.paragraphs:
        content = clean(paragraph.text)
        if not content:
            continue
        if paragraph.style and paragraph.style.name.lower().startswith("heading"):
            if current_paragraphs:
                groups.append((current_title, current_paragraphs))
            current_title, current_paragraphs = content, []
        else:
            current_paragraphs.append(content)
    if current_paragraphs:
        groups.append((current_title, current_paragraphs))
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows[:100]:
            cells = [clean(cell.text) for cell in row.cells[:20]]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            groups.append((f"表格 {table_index}", rows))
    return title, [candidate for index, (heading, lines) in enumerate(groups[:MAX_SECTIONS]) if (candidate := section(heading, lines, index))]


def parse_pdf(blob: bytes) -> tuple[str, list[dict[str, Any]]]:
    pdf = fitz.open(stream=blob, filetype="pdf")
    title = clean(pdf.metadata.get("title"))
    sections: list[dict[str, Any]] = []
    total_chars = 0
    for page_index, page in enumerate(pdf, start=1):
        text = page.get_text("text")
        paragraphs = [clean(line) for line in text.splitlines() if clean(line)]
        total_chars += sum(len(line) for line in paragraphs)
        candidate = section(f"第 {page_index} 页", paragraphs, page_index - 1)
        if candidate:
            sections.append(candidate)
        if len(sections) >= MAX_SECTIONS or total_chars >= MAX_TEXT_CHARS:
            break
    pdf.close()
    return title, sections


def parse_xlsx(blob: bytes) -> tuple[str, list[dict[str, Any]]]:
    workbook = load_workbook(BytesIO(blob), read_only=True, data_only=True)
    sections: list[dict[str, Any]] = []
    total_cells = 0
    for sheet_index, worksheet in enumerate(workbook.worksheets[:MAX_SECTIONS]):
        rows = []
        for row in worksheet.iter_rows(values_only=True):
            values = [clean(value) for value in row[:40]]
            if any(values):
                rows.append(" | ".join(values))
            total_cells += len(values)
            if len(rows) >= MAX_PARAGRAPHS_PER_SECTION or total_cells >= 10_000:
                break
        candidate = section(worksheet.title, rows, sheet_index)
        if candidate:
            sections.append(candidate)
        if total_cells >= 10_000:
            break
    title = clean(workbook.properties.title) or "Excel 工作簿"
    workbook.close()
    return title, sections


def reader_document(title: str, filename: str, sections: list[dict[str, Any]]) -> dict[str, Any]:
    if not sections:
        raise HTTPException(status_code=422, detail="未识别到可阅读的文档正文")
    word_count = sum(len(item["title"]) + sum(len(paragraph) for paragraph in item["paragraphs"]) for item in sections)
    return {
        "title": title[:300] or Path(filename).stem[:300] or "未命名文档",
        "description": sections[0]["paragraphs"][0][:240],
        "sourceUrl": "",
        "siteName": filename[:300] or "上传的文档",
        "fetchedAt": __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat(),
        "wordCount": word_count,
        "engine": "document-processor",
        "sections": sections,
    }


@app.get("/healthz")
async def healthz() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/v1/parse")
async def parse(request: Request) -> dict[str, Any]:
    filename = unquote(request.headers.get("x-source-filename", "").strip())
    extension = Path(filename).suffix.lower()
    if extension not in {".docx", ".md", ".pdf", ".xlsx"}:
        raise HTTPException(status_code=422, detail="仅支持 DOCX、MD、PDF、XLSX")
    blob = await request.body()
    if not blob or len(blob) > MAX_BYTES:
        raise HTTPException(status_code=413, detail="文件大小需在 1 字节到 200 MB 之间")
    try:
        if extension == ".md":
            title, sections = Path(filename).stem or "Markdown 文档", markdown_sections(blob.decode("utf-8-sig", errors="replace"))
        elif extension == ".docx":
            title, sections = parse_docx(blob)
        elif extension == ".pdf":
            title, sections = parse_pdf(blob)
        else:
            title, sections = parse_xlsx(blob)
        fallback_title = Path(filename).stem
        return {"document": reader_document(title or fallback_title, filename, sections)}
    except HTTPException:
        raise
    except Exception as error:
        raise HTTPException(status_code=422, detail=f"文档解析失败：{clean(error)[:300]}") from error
